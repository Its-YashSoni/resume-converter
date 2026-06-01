from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ──
GREEN_HEADER = "92D050"   # Title row + MSP INPUT row
GREEN_DARK   = "375623"
INFO_BG      = "A8D08F"   # Info rows + column headers
YELLOW_SKILL = "FEFF01"   # Skill name cells
YELLOW_TEXT  = "375623"
WHITE        = "FFFFFF"
BLACK        = "000000"


def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, border_color="999999", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_no_border_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_run(para, text, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    run.font.name = font
    return run


def add_bullet(doc, text, bold_part="", normal_part="", size=11, bullet_char="●"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    run0 = p.add_run(bullet_char + "  ")
    run0.font.size = Pt(size)
    run0.font.name = "Calibri"
    run0.font.color.rgb = RGBColor(*hex_to_rgb(BLACK))
    if bold_part:
        add_run(p, bold_part, bold=True, size=size, color=BLACK)
    if normal_part:
        add_run(p, normal_part, bold=False, size=size, color=BLACK)
    if text and not bold_part:
        add_run(p, text, bold=False, size=size, color=BLACK)
    return p


def add_small_bullet(doc, bold_part="", normal_part="", size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run0 = p.add_run("•  ")
    run0.font.size = Pt(size)
    run0.font.name = "Calibri"
    run0.font.color.rgb = RGBColor(*hex_to_rgb(BLACK))
    if bold_part:
        add_run(p, bold_part, bold=True, size=size, color=BLACK)
    if normal_part:
        add_run(p, normal_part, bold=False, size=size, color=BLACK)
    return p


def add_section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK))
    run.underline = True
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ══════════════════════════════════════════════
#  PAGE 1 — Evaluation Table
# ══════════════════════════════════════════════
def build_eval_table(doc, data):
    skills_table_data = data.get('candidate_skills_for_table', [])
    while len(skills_table_data) < 3:
        skills_table_data.append({
            "skill_name": "General Skills",
            "projects_used": "Various Projects",
            "years": "< 1 Yr",
            "description": "Applied across multiple projects"
        })

    total_experience = data.get('total_experience', 'FRESHER') or 'FRESHER'
    col_w = [2.1, 1.1, 3.0, 1.1, 2.06]

    table = doc.add_table(rows=0, cols=5)
    set_no_border_table(table)

    def add_info_row(label, value):
        row = table.add_row()
        c0 = row.cells[0]
        set_cell_bg(c0, INFO_BG)
        set_cell_borders(c0)
        set_col_width(c0, col_w[0])
        p = c0.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, label, bold=True, size=10, color=BLACK, font="Arial")

        merged = row.cells[1]
        for i in range(2, 5):
            merged = merged.merge(row.cells[i])
        set_cell_bg(merged, WHITE)
        set_cell_borders(merged)
        p2 = merged.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        add_run(p2, value,
                bold=(label == "Candidate Name:"),
                size=10 if label != "Candidate Name:" else 11,
                color=BLACK, font="Arial")

    # Row 0: Title — 92D050
    row0 = table.add_row()
    merged_title = row0.cells[0]
    for i in range(1, 5):
        merged_title = merged_title.merge(row0.cells[i])
    set_cell_bg(merged_title, GREEN_HEADER)
    set_cell_borders(merged_title, "777777", "6")
    p = merged_title.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Contractor Connect Skill Evaluation Sheet",
            bold=True, size=11, color=GREEN_DARK, font="Arial")

    # Info rows — A8D08F
    add_info_row("Candidate Name:",        data.get('full_name', '').upper())
    add_info_row("Total Experience:",      total_experience.upper())
    add_info_row("Relevant\nExperience:",  total_experience.upper())
    add_info_row("Notice Period:",         data.get('notice_period', '15 Days'))
    add_info_row("Candidate Aware\nof JD:", "YES")

    # MSP INPUT row — 92D050
    msp_row = table.add_row()
    c0 = msp_row.cells[0]
    set_cell_bg(c0, GREEN_HEADER)
    set_cell_borders(c0)
    set_col_width(c0, col_w[0])
    p = c0.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "o   MSP INPUT", size=10, color=GREEN_DARK, font="Arial")
    merged_msp = msp_row.cells[1]
    for i in range(2, 5):
        merged_msp = merged_msp.merge(msp_row.cells[i])
    set_cell_bg(merged_msp, GREEN_HEADER)
    set_cell_borders(merged_msp)
    p2 = merged_msp.paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, "Supplier Inputs", size=10, color=GREEN_DARK, font="Arial")

    # Column header row — A8D08F
    hdr_row = table.add_row()
    headers = [
        "Candidate Skills",
        "Mandatory/\nOptional",
        "Name of Projects in which the skills were used (add rows if necessary)",
        "No: of years worked in each Project",
        "Description of work done using the skill"
    ]
    for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, headers, col_w)):
        set_cell_bg(cell, INFO_BG)
        set_cell_borders(cell, "777777", "6")
        set_col_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, hdr, bold=True, size=10, color=BLACK, font="Arial")

    # Skill rows — FEFF01
    for skill in skills_table_data[:3]:
        sk_row = table.add_row()
        cells = sk_row.cells

        set_cell_bg(cells[0], YELLOW_SKILL)
        set_cell_borders(cells[0])
        set_col_width(cells[0], col_w[0])
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, skill.get('skill_name', ''), bold=True, size=10, color=YELLOW_TEXT, font="Arial")

        set_cell_bg(cells[1], WHITE)
        set_cell_borders(cells[1])
        set_col_width(cells[1], col_w[1])
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "M", size=10, color=BLACK, font="Arial")

        set_cell_bg(cells[2], WHITE)
        set_cell_borders(cells[2])
        set_col_width(cells[2], col_w[2])
        p = cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, skill.get('projects_used', ''), size=10, color=BLACK, font="Arial")

        set_cell_bg(cells[3], WHITE)
        set_cell_borders(cells[3])
        set_col_width(cells[3], col_w[3])
        p = cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, skill.get('years', '< 1 Yr'), size=10, color=BLACK, font="Arial")

        set_cell_bg(cells[4], WHITE)
        set_cell_borders(cells[4])
        set_col_width(cells[4], col_w[4])
        p = cells[4].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, skill.get('description', ''), size=10, color=BLACK, font="Arial")

    return table


# ══════════════════════════════════════════════
#  PAGES 2-3 — Resume Content
#  ORDER: Summary → Experience → Skills →
#         Projects → Education → Certifications
# ══════════════════════════════════════════════
def build_resume_pages(doc, data):

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, data.get('full_name', '').upper(), bold=True, size=14, color=BLACK)

    # ── 1. PROFILE SUMMARY ──
    summary = data.get('profile_summary', '')
    if summary:
        add_section_label(doc, "PROFILE SUMMARY")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, summary, size=11, color=BLACK)

    # ── 2. WORK EXPERIENCE ──
    experience = data.get('experience', [])
    if experience:
        add_section_label(doc, "WORK EXPERIENCE")
        for exp in experience:
            company     = exp.get('company', '')
            role        = exp.get('role', '')
            duration    = exp.get('duration', '')
            location    = exp.get('location', '')
            company_str = f"{company}, {location}" if location else company

            table = doc.add_table(rows=1, cols=2)
            set_no_border_table(table)
            row = table.rows[0]
            p0 = row.cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(8)
            p0.paragraph_format.space_after = Pt(2)
            add_run(p0, company_str + " | ", bold=True, size=11, color=BLACK)
            add_run(p0, duration, bold=True, italic=True, size=11, color=BLACK)

            if role:
                p_role = doc.add_paragraph()
                p_role.paragraph_format.space_before = Pt(0)
                p_role.paragraph_format.space_after = Pt(4)
                add_run(p_role, role, bold=True, size=11, color=BLACK)

            for bullet_text in exp.get('bullets', []):
                add_bullet(doc, bullet_text, size=11)

    # ── 3. SKILLS ──
    # Each skill item is listed line by line with a ● bullet.
    # Items may come as a comma-separated string inside 'items' field,
    # or as individual entries — we split and bullet each one separately.
    skills = data.get('skills', [])
    if skills:
        add_section_label(doc, "SKILLS")
        for skill in skills:
            category = skill.get('category', '').strip()
            items_raw = skill.get('items', '').strip()

            # If there's a category label, print it as a bold sub-heading
            if category:
                p_cat = doc.add_paragraph()
                p_cat.paragraph_format.space_before = Pt(6)
                p_cat.paragraph_format.space_after = Pt(2)
                add_run(p_cat, category + ":", bold=True, size=11, color=BLACK)

            # Split comma-separated items into individual bullets
            if items_raw:
                individual_items = [i.strip() for i in items_raw.split(',') if i.strip()]
                for item in individual_items:
                    add_bullet(doc, item, size=11)
            elif category:
                # category itself is the skill (no separate items field)
                add_bullet(doc, category, size=11)

    # ── 4. PROJECTS ──
    projects = data.get('projects', [])
    if projects:
        add_section_label(doc, "PROJECTS")
        for proj in projects:
            title   = proj.get('title', '')
            tech    = proj.get('tech_stack', '')
            date    = proj.get('date', '')
            desc    = proj.get('description', '')
            bullets = proj.get('bullets', [])

            table = doc.add_table(rows=1, cols=2)
            set_no_border_table(table)
            row = table.rows[0]
            p0 = row.cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(10)
            p0.paragraph_format.space_after = Pt(2)
            title_text = f"{title}" + (f" ( {tech} )" if tech else "")
            add_run(p0, title_text, bold=True, size=12, color=BLACK)
            p1 = row.cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p1.paragraph_format.space_before = Pt(10)
            p1.paragraph_format.space_after = Pt(2)
            add_run(p1, date, size=10, color=BLACK)

            if desc:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_before = Pt(0)
                p_desc.paragraph_format.space_after = Pt(4)
                add_run(p_desc, desc, italic=True, size=10, color=BLACK)

            for b in bullets:
                if ": " in b and len(b.split(": ")[0]) < 30:
                    parts = b.split(": ", 1)
                    add_small_bullet(doc, bold_part=parts[0] + ": ", normal_part=parts[1], size=10)
                else:
                    add_small_bullet(doc, normal_part=b, size=10)

    # ── 5. EDUCATION ──
    education = data.get('education', [])
    if education:
        add_section_label(doc, "EDUCATION")
        for edu in education:
            inst      = edu.get('institution', '')
            degree    = edu.get('degree', '')
            year      = edu.get('year', '')
            score     = edu.get('score', '')
            full_inst = f"{degree} – {inst}" if degree and inst else (degree or inst)

            table = doc.add_table(rows=1, cols=2)
            set_no_border_table(table)
            row = table.rows[0]

            c0 = row.cells[0]
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after = Pt(2)
            add_run(p0, full_inst, bold=True, size=11, color=BLACK)

            c1 = row.cells[1]
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after = Pt(2)
            add_run(p1, year, size=10, color=BLACK)

            if score:
                p_score = doc.add_paragraph()
                p_score.paragraph_format.space_before = Pt(0)
                p_score.paragraph_format.space_after = Pt(6)
                add_run(p_score, score, size=11, color=BLACK)

    # ── 6. CERTIFICATIONS ──
    certs = data.get('certifications', [])
    if certs:
        add_section_label(doc, "CERTIFICATIONS")
        for cert in certs:
            name   = cert.get('name', '')
            issuer = cert.get('issuer', '')
            desc   = cert.get('description', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_run(p, name, bold=True, size=11, color=BLACK)
            if issuer:
                add_run(p, f" – {issuer}", italic=True, size=11, color=BLACK)
            if desc:
                add_bullet(doc, desc, size=11)

    # ── TRAINING ──
    training = data.get('training', [])
    if training:
        add_section_label(doc, "TRAINING")
        for t in training:
            name = t if isinstance(t, str) else t.get('name', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_run(p, name, bold=True, size=11, color=BLACK)

    # ── CO-CURRICULUM ──
    co = data.get('co_curriculum', [])
    if co:
        add_section_label(doc, "CO CURRICULUM")
        for item in co:
            add_bullet(doc, item if isinstance(item, str) else str(item), size=11)

    # ── EXTRA CURRICULUM ──
    extra = data.get('extra_curriculum', [])
    if extra:
        add_section_label(doc, "EXTRA CURRICULUM")
        for item in extra:
            add_bullet(doc, item if isinstance(item, str) else str(item), size=11)

    # ── SOFT SKILLS ──
    soft = data.get('soft_skills', [])
    if soft:
        add_section_label(doc, "SOFT SKILLS")
        for item in soft:
            add_bullet(doc, item if isinstance(item, str) else str(item), size=11)

    # ── LANGUAGES ──
    langs = data.get('languages', [])
    if langs:
        add_section_label(doc, "LANGUAGES")
        for lang in langs:
            add_bullet(doc, lang if isinstance(lang, str) else str(lang), size=11)


# ══════════════════════════════════════════════
#  MAIN BUILDER
# ══════════════════════════════════════════════
def build_resume_docx(data, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    build_eval_table(doc, data)
    add_page_break(doc)
    build_resume_pages(doc, data)

    doc.save(output_path)
    return output_path
